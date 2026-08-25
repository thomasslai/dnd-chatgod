"""Build a lore-only vector index.

By default this script indexes documents from data/sorted_data/lore and writes
to vector_db. Campaign-memory sources are intentionally excluded and belong in
the SQLite pipeline (ingest_campaign.py).
"""

from __future__ import annotations

import os
import time
from pathlib import Path

import PyPDF2
from dotenv import load_dotenv
from docx import Document as DocxDocument
from llama_index.core import Document, VectorStoreIndex
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.embeddings.openai import OpenAIEmbedding
from odf import teletype
from odf.opendocument import load as load_odt
from odf.text import P

from settings import (
	CHAT_EMBED_MODEL,
	LORE_CAMPAIGN_PATH_MARKERS,
	LORE_CHUNK_OVERLAP,
	LORE_CHUNK_SIZE,
	LORE_DEFAULT_FOLDER,
	LORE_PERSIST_DIR,
	LORE_SUPPORTED_EXTENSIONS,
)

load_dotenv()

ROOT = Path(__file__).resolve().parent
DEFAULT_LORE_FOLDER = LORE_DEFAULT_FOLDER
DEFAULT_PERSIST_DIR = LORE_PERSIST_DIR
CAMPAIGN_PATH_MARKERS = LORE_CAMPAIGN_PATH_MARKERS
SUPPORTED_EXTENSIONS = LORE_SUPPORTED_EXTENSIONS


def resolve_lore_folder() -> Path:
	"""Resolve the lore input folder from env vars with a safe default."""
	env_value = os.getenv("LORE_DOCS_FOLDER")
	folder = Path(env_value).expanduser() if env_value else DEFAULT_LORE_FOLDER
	if not folder.is_absolute():
		folder = (ROOT / folder).resolve()
	return folder


def ensure_lore_only_path(folder: Path) -> None:
	"""Guard against accidentally indexing campaign-memory sources."""
	lower_parts = {part.lower() for part in folder.parts}
	if CAMPAIGN_PATH_MARKERS & lower_parts:
		markers = ", ".join(sorted(CAMPAIGN_PATH_MARKERS))
		raise ValueError(
			f"Refusing to index campaign-memory path '{folder}'. "
			f"Path contains one of: {markers}."
		)


def _extract_text(file_path: Path) -> str:
	"""Extract plain text from supported lore file formats."""
	suffix = file_path.suffix.lower()

	if suffix == ".txt":
		return file_path.read_text(encoding="utf-8", errors="ignore")

	if suffix == ".docx":
		doc = DocxDocument(str(file_path))
		return "\n".join(paragraph.text for paragraph in doc.paragraphs)

	if suffix == ".pdf":
		text_parts: list[str] = []
		with file_path.open("rb") as handle:
			reader = PyPDF2.PdfReader(handle)
			for page in reader.pages:
				text_parts.append(page.extract_text() or "")
		return "\n".join(text_parts)

	if suffix == ".odt":
		odt_doc = load_odt(str(file_path))
		paragraphs = [teletype.extractText(elem) for elem in odt_doc.getElementsByType(P)]
		return "\n".join(paragraphs)

	return ""


def load_lore_documents(docs_folder: Path) -> list[Document]:
	"""Load lore documents with explicit parsing to avoid binary gibberish."""
	documents: list[Document] = []
	for file_path in sorted(docs_folder.rglob("*")):
		if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
			continue

		text = _extract_text(file_path).strip()
		if not text:
			continue

		relative_path = file_path.relative_to(docs_folder).as_posix()
		documents.append(
			Document(
				text=text,
				metadata={
					"file_name": file_path.name,
					"file_path": relative_path,
				},
			)
		)
	return documents


def main() -> None:
	started_at = time.time()
	docs_folder = resolve_lore_folder()
	persist_dir = DEFAULT_PERSIST_DIR

	if not docs_folder.exists():
		raise FileNotFoundError(f"Lore docs folder not found: {docs_folder}")
	ensure_lore_only_path(docs_folder)

	documents = load_lore_documents(docs_folder)
	if not documents:
		raise RuntimeError(f"No parseable lore documents found in: {docs_folder}")

	node_parser = SimpleNodeParser.from_defaults(chunk_size=LORE_CHUNK_SIZE, chunk_overlap=LORE_CHUNK_OVERLAP)
	nodes = node_parser.get_nodes_from_documents(documents)

	api_key = os.getenv("OPENAI_API_KEY")
	embed_model = OpenAIEmbedding(model=CHAT_EMBED_MODEL, api_key=api_key)

	index = VectorStoreIndex(nodes, embed_model=embed_model)
	index.storage_context.persist(persist_dir=str(persist_dir))

	elapsed = time.time() - started_at
	print(f"Indexed {len(documents)} lore docs into '{persist_dir}' in {elapsed:.2f}s")


if __name__ == "__main__":
	main()